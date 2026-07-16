from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "design" / "ai_trade_signal_execution_design_document_v0.2.docx"
ASSET = ROOT / "docs" / "design" / "assets" / "signal_execution_lifecycle.png"

NAVY = "17365D"
BLUE = "2E74B5"
INK = "1F2937"
MUTED = "667085"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_widths(table, widths_inches):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            set_cell_margin(cell)


def font(size, bold=False, color=INK, italic=False):
    return {"size": Pt(size), "bold": bold, "color": RGBColor.from_string(color), "italic": italic}


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    for key, value in kwargs.items():
        if key == "color":
            run.font.color.rgb = value
        elif key == "size":
            run.font.size = value
        elif key == "bold":
            run.bold = value
        elif key == "italic":
            run.italic = value
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    return run


def add_text(doc, text, style=None, after=7, before=0, color=INK, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    add_run(p, text, **font(size, bold=bold, color=color, italic=italic))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text, **font(10.3))
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, f"{label}: ", **font(10.5, bold=True, color=NAVY))
    add_run(p, text, **font(10.5, color=INK))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def draw_lifecycle(path):
    width, height = 1600, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        regular = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 27)
        bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 29)
        small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 21)
        title = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 38)
    except OSError:
        regular = bold = small = title = ImageFont.load_default()

    draw.text((70, 40), "Phase 1 to Phase 2: decision-to-trade lifecycle", font=title, fill="#17365D")
    draw.text((70, 94), "Macro narrows the universe. The signal system decides whether, when, and how to trade.", font=regular, fill="#667085")

    steps = [
        ("1", "MACRO FILTER", "regime, sectors,\nshortlist", "EAF2F8"),
        ("2", "SETUP", "technical rules,\nentry context", "EAF2F8"),
        ("3", "SIGNAL", "entry, stop, target,\nexpiry", "FFF4D6"),
        ("4", "RISK GATE", "size, limits,\napproval", "FDECEC"),
        ("5", "EXECUTE", "venue order,\nfill tracking", "EAF2F8"),
        ("6", "TRACK + LEARN", "monitor thesis,\nmeasure outcome", "EAF2F8"),
    ]
    x0, y, w, h, gap = 65, 230, 220, 205, 35
    for idx, (num, heading, body, fill) in enumerate(steps):
        x = x0 + idx * (w + gap)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=f"#{fill}", outline="#B7C6D8", width=2)
        draw.ellipse((x + 18, y + 17, x + 58, y + 57), fill="#17365D")
        draw.text((x + 31, y + 22), num, font=small, fill="white")
        draw.text((x + 18, y + 79), heading, font=bold, fill="#17365D")
        draw.multiline_text((x + 18, y + 124), body, font=regular, fill="#1F2937", spacing=5)
        if idx < len(steps) - 1:
            ax = x + w + 7
            draw.line((ax, y + h / 2, ax + gap - 14, y + h / 2), fill="#7087A3", width=4)
            draw.polygon([(ax + gap - 14, y + h / 2), (ax + gap - 28, y + h / 2 - 9), (ax + gap - 28, y + h / 2 + 9)], fill="#7087A3")

    draw.line((1335, 510, 1335, 615), fill="#7087A3", width=4)
    draw.line((1335, 615, 174, 615), fill="#7087A3", width=4)
    draw.line((174, 615, 174, 470), fill="#7087A3", width=4)
    draw.polygon([(174, 470), (165, 486), (183, 486)], fill="#7087A3")
    draw.rounded_rectangle((350, 566, 1250, 682), radius=18, fill="#F5F7FA", outline="#D0D9E5", width=2)
    draw.text((385, 586), "Feedback loop: compare expected and actual fills and P&L.", font=regular, fill="#334155")
    draw.text((430, 622), "Revise or retire the strategy version; no direct LLM-to-order path.", font=bold, fill="#9B1C1C")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "AI TRADE | SIGNAL & EXECUTION DESIGN", **font(8.5, bold=True, color=MUTED))
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Living design document | v0.1 | Internal working draft", **font(8.5, color=MUTED))


def build_doc():
    draw_lifecycle(ASSET)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in (("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, NAVY)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(15 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "AI TRADE", **font(11, bold=True, color=BLUE))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Signal & Execution Design Document", **font(26, bold=True, color=NAVY))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "Phase 2 - turning a macro shortlist into a controlled trade lifecycle | Version 0.2", **font(13, color=MUTED))

    meta = doc.add_table(rows=3, cols=2)
    set_table_widths(meta, [1.55, 4.95])
    meta_data = [
        ("Document status", "Living design document. Strategy 01 is defined for implementation and backtest; it is not live-trading approval."),
        ("Relationship to Phase 1", "Phase 1 provides macro context and a shortlist; this document defines the technical signal, risk, execution, and learning layer."),
        ("Current scope", "Strategy 01 on SPY through IBKR paper. Backtest and shadow mode precede all order submission."),
    ]
    for row, (label, value) in zip(meta.rows, meta_data):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PALE_BLUE)
        add_run(row.cells[0].paragraphs[0], label, **font(9.7, bold=True, color=NAVY))
        add_run(row.cells[1].paragraphs[0], value, **font(9.7))
    doc.add_paragraph()

    add_callout(doc, "Core design decision", "Phase 2 is not separate from Phase 1; it is the next decision layer. Macro tells us which environments and candidates deserve attention. Phase 2 turns a candidate into a testable trade plan and only then, after risk checks, an executable order.", PALE_GOLD)

    add_text(doc, "Contents", style="Heading 1", after=4)
    for item in [
        "1. From macro shortlist to trade lifecycle",
        "2. Decision layers and handoff rules",
        "3. Signal contract: what a valid trade must contain",
        "4. Operating modes and safety gates",
        "5. Initial design decisions and next step",
        "6. Strategy 01 - Bill Williams Alligator + Heikin Ashi + RRMS",
    ]:
        add_bullet(doc, item)

    add_text(doc, "1. From Macro Shortlist to Trade Lifecycle", style="Heading 1")
    add_text(doc, "Purpose of this chapter", style="Heading 2")
    add_text(doc, "This chapter establishes the first design boundary for the AI Trade project. It explains how the Phase 1 Macro Dashboard and Phase 2 trading system work together, what each layer is responsible for, and how a trade should travel from idea to measurement. It is intentionally a design document, not yet a product specification or live-trading authorization.")

    add_text(doc, "The handoff", style="Heading 2")
    add_text(doc, "Phase 1 has already narrowed the opportunity space: the macro regime, preferred sectors, risk appetite, and a shortlist of US stocks provide context. That work is necessary, but it does not yet answer the trading questions that determine whether capital should be put at risk today.")
    add_text(doc, "Phase 2 begins when a shortlisted asset is assessed by an explicit strategy. The strategy needs technical entry and exit rules, a defined invalidation point, position-sizing logic, and an objective way to evaluate the result after the trade closes.")

    add_text(doc, "What each phase owns", style="Heading 2")
    handoff = doc.add_table(rows=1, cols=3)
    set_table_widths(handoff, [1.35, 2.48, 2.67])
    headers = ["Layer", "Primary question", "Output"]
    for cell, value in zip(handoff.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.6, bold=True, color=WHITE))
    set_repeat_table_header(handoff.rows[0])
    rows = [
        ("Phase 1: Macro", "Where should attention go?", "Regime, risk stance, preferred sectors/themes, and a candidate shortlist."),
        ("Phase 2A: Setup", "Is there a tradable technical setup?", "A rule-based setup with market, timeframe, conditions, and expiry."),
        ("Phase 2B: Signal", "Should we enter now, and on what terms?", "Entry, stop, target/exit logic, expected reward-to-risk, and confidence."),
        ("Phase 2C: Risk + execution", "May this intent become an order?", "Approved/rejected intent, permitted size, venue order, and verified fill."),
        ("Track + improve", "Did reality match the thesis and simulation?", "Trade record, post-trade review, strategy scorecard, and version decision."),
    ]
    for values in rows:
        cells = handoff.add_row().cells
        for cell, value in zip(cells, values):
            add_run(cell.paragraphs[0], value, **font(9.2))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()

    add_text(doc, "Lifecycle visual", style="Heading 2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSET), width=Inches(6.55))
    p.paragraph_format.space_after = Pt(2)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(caption, "Figure 1. The strategy loop is closed only when real outcomes feed back into the next version of the strategy.", **font(8.5, italic=True, color=MUTED))

    add_text(doc, "2. Decision Layers and Handoff Rules", style="Heading 1")
    add_text(doc, "The system should preserve a clear separation between deciding, approving, and doing. This prevents a promising chart or an AI-generated explanation from bypassing risk control.")
    layers = [
        ("1. Macro context", "Optional filter: market regime, sector preference, and event risk determine whether a setup is allowed or should be treated cautiously."),
        ("2. Strategy setup", "Deterministic technical criteria identify the market, timeframe, and conditions that qualify a candidate for a possible trade."),
        ("3. Signal", "A qualifying setup becomes an intent only when the defined trigger occurs. The signal contains no discretionary gaps."),
        ("4. Risk gateway", "Checks permitted market, capital allocation, maximum loss, leverage, existing exposure, trading hours, and daily stop rules."),
        ("5. Execution adapter", "Maps an approved intent to venue-specific order instructions and records acknowledgements, partial fills, rejections, and cancellations."),
        ("6. Tracking and review", "Reconciles broker truth with internal records and measures the result against the strategy's expected behavior."),
    ]
    for heading, body in layers:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run(p, f"{heading}. ", **font(10.5, bold=True, color=NAVY))
        add_run(p, body, **font(10.5))

    add_text(doc, "3. Signal Contract: What a Valid Trade Must Contain", style="Heading 1")
    add_text(doc, "A signal is a structured trade proposal, not a vague direction such as 'buy Tesla.' Before it can reach the risk gateway, it must provide the following fields.")
    contract = doc.add_table(rows=1, cols=2)
    set_table_widths(contract, [1.8, 4.7])
    for cell, value in zip(contract.rows[0].cells, ["Field", "Required meaning"]):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.6, bold=True, color=WHITE))
    set_repeat_table_header(contract.rows[0])
    fields = [
        ("Identity", "Strategy version, venue, symbol, timeframe, and a unique signal ID."),
        ("Why now", "Exact setup and trigger conditions that fired; macro context may be attached as a filter, not invented after entry."),
        ("Entry", "Order side, allowed entry price/range, order type, and time-to-live."),
        ("Invalidation", "Stop-loss or condition that proves the thesis wrong and ends the trade."),
        ("Exit plan", "Profit target, trailing or time-based exit, and rules for partial exits if used."),
        ("Risk", "Maximum loss, proposed position size, leverage if applicable, and expected fees/funding assumptions."),
        ("Audit", "Data timestamp, model/configuration version, and every later state transition."),
    ]
    for field, meaning in fields:
        cells = contract.add_row().cells
        add_run(cells[0].paragraphs[0], field, **font(9.3, bold=True, color=NAVY))
        add_run(cells[1].paragraphs[0], meaning, **font(9.3))
    doc.add_paragraph()

    add_callout(doc, "Non-negotiable rule", "An LLM can help explain, test ideas, or summarize context. It does not submit an order. Only an explicit strategy intent that passes the risk gateway may reach an execution adapter.", PALE_RED)

    add_text(doc, "4. Operating Modes and Safety Gates", style="Heading 1")
    add_text(doc, "The same strategy should advance through explicit modes. Promotion is deliberate; it never happens simply because a backtest looks attractive.")
    mode_table = doc.add_table(rows=1, cols=3)
    set_table_widths(mode_table, [1.1, 2.2, 3.2])
    for cell, value in zip(mode_table.rows[0].cells, ["Mode", "What runs", "Gate to leave the mode"]):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.6, bold=True, color=WHITE))
    set_repeat_table_header(mode_table.rows[0])
    modes = [
        ("Backtest", "Historical data and simulated orders.", "Reproducible report with realistic costs and an untouched out-of-sample period."),
        ("Shadow", "Live data; hypothetical trades only.", "Signals, timing, and health checks behave as specified for a meaningful observation period."),
        ("Paper", "Orders sent only to a paper environment.", "Order states, restarts, fills, and balances reconcile with the venue."),
        ("Limited live", "Small, constrained real orders.", "Explicit user authorization, limits configured, alerts working, and manual kill switch tested."),
    ]
    for mode, activity, gate in modes:
        cells = mode_table.add_row().cells
        for cell, text in zip(cells, (mode, activity, gate)):
            add_run(cell.paragraphs[0], text, **font(9.2, bold=(cell == cells[0]), color=NAVY if cell == cells[0] else INK))
    doc.add_paragraph()

    add_text(doc, "5. Initial Design Decisions and Next Step", style="Heading 1")
    add_text(doc, "Decisions made in this chapter", style="Heading 2")
    for item in [
        "Phase 1 and Phase 2 overlap by design: the macro layer filters the universe; it does not replace a technical strategy.",
        "The current project owns signal generation, backtesting, risk controls, broker/exchange adapters, reconciliation, monitoring, and strategy improvement.",
        "The initial release focuses on a single small experiment rather than a multi-strategy dashboard or unrestricted live trading.",
        "Broker/exchange records are the source of truth for positions, balances, orders, and fills.",
        "Any strategy change creates a new version that must repeat backtest, shadow, and paper validation before live use.",
    ]:
        add_bullet(doc, item)

    add_text(doc, "Next design step", style="Heading 2")
    add_text(doc, "Define the first experiment in a one-page strategy brief. The brief must choose one venue, one market, one timeframe, one fully rule-based setup, and the costs/limits that a realistic backtest must include. That brief becomes the input to the data-collection and backtesting work.")
    add_callout(doc, "First experiment selected", "Strategy 01 will be developed on SPY with IBKR paper trading. MEXC spot and futures remain later extensions after the core cycle has been proven.", PALE_GOLD)

    add_text(doc, "6. Strategy 01 - Bill Williams Alligator + Heikin Ashi + RRMS", style="Heading 1")
    add_callout(doc, "Implementation status", "This is the first strategy selected for implementation. Its first use is a historical backtest, then shadow mode, then IBKR paper trading. It is not approved for live trading.", PALE_GOLD)

    add_text(doc, "6.1 Strategy objective and scope", style="Heading 2")
    add_text(doc, "Strategy 01 is an intraday trend-following system for the S&P 500 using the SPY ETF as the initial executable instrument. The Macro Dashboard is the context layer: it determines whether S&P 500 trades are allowed. The 1-hour chart confirms trend direction, and the 15-minute chart determines the setup and entry.")
    scope = doc.add_table(rows=1, cols=2)
    set_table_widths(scope, [1.85, 4.65])
    for cell, value in zip(scope.rows[0].cells, ["Design item", "Decision"]):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.6, bold=True, color=WHITE))
    set_repeat_table_header(scope.rows[0])
    scope_rows = [
        ("Strategy name", "Strategy 01 - Bill Williams Alligator + Heikin Ashi + RRMS"),
        ("Market", "S&P 500, initially executed as SPY through IBKR paper trading. A continuous S&P chart may be used for research only; it is not an executable contract."),
        ("Macro filter", "S&P 500 trading is allowed only when the Phase 1 macro stance permits it."),
        ("Trend confirmation", "1-hour chart"),
        ("Entry timeframe", "15-minute chart"),
        ("Direction", "Long and short rules are defined symmetrically; the initial backtest will report their results separately."),
        ("Initial reward model", "1R target: target distance equals the actual entry-to-stop distance. This is provisional and must be tested."),
    ]
    for label, value in scope_rows:
        cells = scope.add_row().cells
        add_run(cells[0].paragraphs[0], label, **font(9.3, bold=True, color=NAVY))
        add_run(cells[1].paragraphs[0], value, **font(9.3))
    doc.add_paragraph()

    add_text(doc, "6.2 Indicator definitions and no-lookahead rule", style="Heading 2")
    add_text(doc, "The Alligator uses three smoothed moving averages of median price, where median price is (high + low) / 2. The Jaw is a 13-period smoothed moving average displayed 8 bars forward; Teeth is an 8-period average displayed 5 bars forward; Lips is a 5-period average displayed 3 bars forward. Heikin Ashi candles are used to reduce visual noise in the entry signal.")
    add_callout(doc, "Backtest integrity", "Forward offsets are display offsets only. The implementation must never use candles that did not exist at the decision time. Every 1-hour and 15-minute rule is evaluated only from completed data available at that time.", PALE_RED)

    add_text(doc, "6.3 Trade rules", style="Heading 2")
    rules = doc.add_table(rows=1, cols=3)
    set_table_widths(rules, [1.2, 2.65, 2.65])
    for cell, value in zip(rules.rows[0].cells, ["Rule", "Long", "Short"]):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.6, bold=True, color=WHITE))
    set_repeat_table_header(rules.rows[0])
    rule_rows = [
        ("Macro", "Macro stance permits the trade.", "Macro stance permits the trade."),
        ("1-hour confirmation", "Lips above Teeth above Jaw; all three rise; mouth is open rather than compressed.", "Lips below Teeth below Jaw; all three fall; mouth is open rather than compressed."),
        ("15-minute setup", "Same bullish line order and open-mouth condition.", "Same bearish line order and open-mouth condition."),
        ("Trigger", "Completed Heikin Ashi candle body closes above Lips.", "Completed Heikin Ashi candle body closes below Lips."),
        ("Entry", "Next 15-minute candle open after the trigger.", "Next 15-minute candle open after the trigger."),
        ("Stop", "Below the 15-minute Jaw plus a small configured buffer.", "Above the 15-minute Jaw plus a small configured buffer."),
        ("Target", "Entry + 1R, where R is the actual entry-to-stop distance.", "Entry - 1R, where R is the actual entry-to-stop distance."),
        ("Trend exit", "A confirmed opposite setup or an Alligator mouth-closing rule, to be tested alongside the fixed target.", "A confirmed opposite setup or an Alligator mouth-closing rule, to be tested alongside the fixed target."),
    ]
    for values in rule_rows:
        cells = rules.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            add_run(cell.paragraphs[0], value, **font(8.9, bold=(index == 0), color=NAVY if index == 0 else INK))
    doc.add_paragraph()
    add_callout(doc, "Parameter still to formalize", "'Mouth open' and 'parallel' are visual concepts. Before coding, the strategy brief will set numerical thresholds for line separation, slope, and whether separation must be widening. No discretionary chart reading will be used in the backtest.", PALE_GOLD)

    add_text(doc, "6.4 RRMS position-sizing overlay", style="Heading 2")
    add_text(doc, "RRMS is separate from the entry logic. The Alligator and Heikin Ashi rules determine whether a trade exists and the Jaw determines its stop. RRMS determines the maximum account loss and converts that loss budget into position size.")
    rrms = doc.add_table(rows=1, cols=4)
    set_table_widths(rrms, [1.5, 1.55, 1.8, 1.65])
    for cell, value in zip(rrms.rows[0].cells, ["State", "Risk of account equity", "Purpose", "If stop-loss hits"]):
        set_cell_shading(cell, NAVY)
        add_run(cell.paragraphs[0], value, **font(9.2, bold=True, color=WHITE))
    set_repeat_table_header(rrms.rows[0])
    rrms_rows = [
        ("Normal", "0.15%", "First trade / reset tier", "Move to Recovery 1"),
        ("Recovery 1", "0.35%", "After one stop-loss", "Move to Recovery 2"),
        ("Recovery 2", "0.70%", "After two consecutive stop-losses", "Move to Recovery 3"),
        ("Recovery 3", "1.50%", "After three consecutive stop-losses", "Stop trading and require review"),
    ]
    for values in rrms_rows:
        cells = rrms.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            add_run(cell.paragraphs[0], value, **font(8.9, bold=(index == 0), color=NAVY if index == 0 else INK))
    doc.add_paragraph()
    add_text(doc, "Sizing formula", style="Heading 3", after=4)
    formula = doc.add_table(rows=1, cols=1)
    set_table_widths(formula, [6.5])
    set_cell_shading(formula.cell(0, 0), PALE_BLUE)
    p = formula.cell(0, 0).paragraphs[0]
    add_run(p, "Risk dollars = current account equity x RRMS tier percentage\n", **font(10.2, bold=True, color=NAVY))
    add_run(p, "Risk per unit = absolute value of (planned entry price - stop price)\n", **font(10.2, bold=True, color=NAVY))
    add_run(p, "Quantity = floor(risk dollars / risk per unit), limited by available buying power and venue minimum size.", **font(10.2, bold=True, color=NAVY))
    doc.add_paragraph()
    add_bullet(doc, "A profitable closed trade resets RRMS to the Normal 0.15% tier.")
    add_bullet(doc, "RRMS advances only after a stop-loss event. A break-even or other discretionary exit does not advance the tier and is flagged for review.")
    add_bullet(doc, "After four consecutive stop-losses, the system blocks new Strategy 01 entries until a manual review. It does not continue increasing position size.")
    add_bullet(doc, "The backtest will report strategy results with fixed sizing and with RRMS separately. RRMS must not hide whether the entry strategy itself has an edge.")

    add_text(doc, "6.5 Implementation sequence", style="Heading 2")
    for item in [
        "Collect and validate SPY 1-hour and 15-minute historical bars.",
        "Implement Heikin Ashi and no-lookahead Alligator calculations.",
        "Define numerical open-mouth, parallel-slope, Jaw-buffer, and trading-hours parameters.",
        "Backtest the entry strategy with fixed sizing first, then apply the RRMS overlay as a separate result series.",
        "Run shadow mode on live data before IBKR paper orders are considered.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "AI Trade Signal & Execution Design Document"
    doc.core_properties.subject = "Phase 2 signal and execution architecture"
    doc.core_properties.author = "AI Trade project"
    doc.core_properties.comments = "Living design document, version 0.2"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
