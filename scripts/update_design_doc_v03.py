"""Create Design Document v0.3 with the trading-automation maturity model."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "design" / "ai_trade_signal_execution_design_document_v0.2.docx"
DESTINATION = ROOT / "docs" / "design" / "ai_trade_signal_execution_design_document_v0.3.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, value: str, bold: bool = False, color: str | None = None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def main() -> None:
    document = Document(SOURCE)

    # Maintain the existing document's title block, only advancing the version.
    for paragraph in document.paragraphs:
        if "Version 0.2" in paragraph.text:
            for run in paragraph.runs:
                if "Version 0.2" in run.text:
                    run.text = run.text.replace("Version 0.2", "Version 0.3")
            break

    # Update the manually maintained contents list.
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "6. Strategy 01 - Bill Williams Alligator + Heikin Ashi + RRMS":
            added = document.add_paragraph("7. Trading automation maturity model", style="List Bullet")
            paragraph._p.addnext(added._p)
            break

    document.add_heading("7. Trading Automation Maturity Model", level=1)
    document.add_paragraph(
        "Algo trading is the use of programmed rules to make trading decisions and/or execute them. "
        "Order types such as market, limit, stop-loss, and take-profit are broker instructions; "
        "the algorithm is the decision and risk logic that determines whether, when, and how those instructions are used."
    )

    document.add_heading("7.1 Four levels of automation", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = (Inches(1.2), Inches(1.8), Inches(3.5))
    headers = ("Level", "Decision owner", "SPY Alligator example")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        set_cell_shading(cell, "1F4E78")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    rows = (
        ("1. Order automation", "Trader", "The trader identifies a setup. A position sizer calculates quantity from permitted account risk and the selected stop, then prepares or submits an entry with its protective stop and 1R target."),
        ("2. Signal automation", "Strategy + trader approval", "The system detects the 1h, 15m, and 5m alignment, applies time and risk gates, and records or alerts a complete proposed trade. The trader decides whether to approve it."),
        ("3. Rule-based execution", "Strategy", "After the same checks, the system independently sends and manages the bracket order, records fills and errors, and applies defined exit rules. Human intervention is still available through a kill switch."),
        ("4. Systematic portfolio system", "Portfolio-level rules", "Several validated strategies and instruments share capital. The system controls total risk and correlation, monitors real-versus-expected performance, pauses degradation, and promotes only validated strategy versions."),
    )
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].width = widths[index]
            set_cell_text(cells[index], value)

    document.add_heading("7.2 Position Sizer classification", level=2)
    document.add_paragraph(
        "The MetaTrader Position Sizer used in the course is Level 1 order automation when it takes the trader's chosen entry and stop, calculates position size from a risk amount or percentage, derives the take-profit from the chosen reward/risk ratio, and creates the ready-to-send order. The trader still chooses the trade. If a tool only displays the calculation and does not create an order ticket, it is a risk calculator rather than order automation."
    )
    document.add_paragraph(
        "For this project, the first safe execution milestone is a broker-neutral position-sizing and bracket-order preview. It must remain manual-confirmation only. This is distinct from and comes before any automated order submission."
    )

    document.add_heading("7.3 Current project position and promotion path", level=2)
    document.add_paragraph(
        "The project currently has components of Level 2: read-only IBKR data, locally cached historical bars, deterministic Alligator strategies, backtests, saved reports/visuals, and a time-gated shadow-signal process. It intentionally has no broker-order permission."
    )
    add_bullet(document, "Next: add a manual position-sizing and bracket preview; no order is transmitted until the trader explicitly confirms it.")
    add_bullet(document, "Then: connect the preview to IBKR paper trading, with reconciliation, protective-order handling, daily loss limits, alerts, and a manual kill switch.")
    add_bullet(document, "Later: consider rule-based execution only after a strategy has passed long out-of-sample tests, realistic cost/slippage tests, Monte Carlo stress tests, and sustained shadow and paper evidence.")
    add_bullet(document, "A multi-strategy portfolio system is a future architecture stage, not a prerequisite for validating Strategy 01.")

    document.save(DESTINATION)
    print(DESTINATION)


if __name__ == "__main__":
    main()
